"""Create Bill"""
import docx
from docx import Document
from docx.shared import Pt
import pandas as pd
import openpyxl 
def create_bill():
    """Create Bill"""
    table = pd.read_excel('Sales.xlsx',engine='openpyxl')
    name = pd.read_excel('Sales.xlsx',engine='openpyxl')['Name']
    add_1 = table['Address1_1']
    add_2 = table['Address1_2']
    rc = table['Receipt_ID']
    date = table['Date']
    product = table['Product']
    unit = table['Unit']
    unit_price = table['Unit_Price']

    doc = Document("RC.docx")
    for i in range(len(table)):
        new_data = []
        doc.tables[1].cell(0,0).paragraphs[1].text = name[i]
        new_data.append(doc.tables[1].cell(0,0).paragraphs[1])
        doc.tables[1].cell(0,0).paragraphs[2].text = add_1[i]
        new_data.append(doc.tables[1].cell(0,0).paragraphs[2])
        doc.tables[1].cell(0,0).paragraphs[3].text = add_2[i]
        new_data.append(doc.tables[1].cell(0,0).paragraphs[3])
        doc.tables[1].cell(0,0).paragraphs[4].text = 'TAX ID: ' + str(table['Tax_ID'][i])
        new_data.append(doc.tables[1].cell(0,0).paragraphs[4])
        doc.tables[1].cell(0,2).paragraphs[0].text = rc[i]
        new_data.append(doc.tables[1].cell(0,2).paragraphs[0])
        doc.tables[1].cell(0,2).paragraphs[1].text = str(date[i]).split()[0]
        new_data.append(doc.tables[1].cell(0,2).paragraphs[1])
        doc.tables[2].cell(1,0).paragraphs[0].text = product[i]
        new_data.append(doc.tables[2].cell(1,0).paragraphs[0])
        doc.tables[2].cell(1,1).paragraphs[0].text = str(unit[i])
        new_data.append(doc.tables[2].cell(1,1).paragraphs[0])
        doc.tables[2].cell(1,2).paragraphs[0].text = str(unit_price[i])
        new_data.append(doc.tables[2].cell(1,2).paragraphs[0])
        doc.tables[2].cell(1,3).paragraphs[0].text = str(unit_price[i]*unit[i])
        new_data.append(doc.tables[2].cell(1,3).paragraphs[0])
        doc.tables[3].cell(0,1).paragraphs[0].text = str(unit_price[i]*unit[i]) + ' THB'
        new_data.append(doc.tables[3].cell(0,1).paragraphs[0])
        doc.tables[3].cell(0,1).paragraphs[1].text = str(int(unit_price[i]*unit[i]*0.07)) + ' THB'
        new_data.append(doc.tables[3].cell(0,1).paragraphs[1])
        doc.tables[3].cell(0,1).paragraphs[2].text = str(int(unit_price[i]*unit[i]*1.07)) + ' THB'
        new_data.append(doc.tables[3].cell(0,1).paragraphs[2])
        for data in new_data:
            data.runs[0].font.name = 'Cordia New (Body CS)'
            data.runs[0].font.size = Pt(14)
        doc.save(rc[i]+'.docx')
create_bill()
